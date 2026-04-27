"""
Zakładka Konwerter PDF → Word z OCR
Rozpoznaje tekst nawet na zeskanowanych PDF (obrazach)
"""
from core.engine_manager import PDFEngineManager
from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QPushButton,
    QLabel, QFileDialog, QMessageBox, QProgressBar,
    QGroupBox, QSpinBox, QComboBox,
    QTextEdit, QRadioButton, QApplication
)
from PyQt5.QtCore import QThread, pyqtSignal, Qt
import os
import tempfile
import fitz
import time

try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    import docx
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False


class OCRConvertThread(QThread):
    finished = pyqtSignal(bool, str)
    progress = pyqtSignal(int)
    status = pyqtSignal(str)
    log = pyqtSignal(str)

    def __init__(self, pdf_path, output_path, mode="ocr", dpi=300, lang="pol+eng"):
        super().__init__()
        self.pdf_path = pdf_path
        self.output_path = output_path
        self.mode = mode
        self.dpi = dpi
        self.lang = lang

    def run(self):
        try:
            has_text = self.check_if_pdf_has_text()
            
            if has_text and self.mode == "pdf2docx":
                self.log.emit("📄 Wykryto tekstowy PDF - używam pdf2docx...")
                self.convert_with_pdf2docx()
            else:
                self.log.emit("🖼️ Wykryto zeskanowany PDF (obrazy) - używam OCR...")
                self.convert_with_ocr()
        except Exception as e:
            self.finished.emit(False, f"Błąd: {str(e)}")

    def check_if_pdf_has_text(self):
        """Sprawdza czy PDF zawiera tekst (nie jest zeskanowany)"""
        try:
            # Użyj wspólnego silnika lub otwórz bezpośrednio
            engine = PDFEngineManager()
            doc = engine.open_pdf(self.pdf_path)
            
            text_found = False
            for i, page in enumerate(doc):
                if i >= 3:  # Sprawdź tylko pierwsze 3 strony
                    break
                text = page.get_text()
                if len(text.strip()) > 50:  # Znaleziono znaczący tekst
                    text_found = True
                    break
            return text_found
        except Exception as e:
            print(f"Błąd sprawdzania tekstu: {e}")
            return False

    def convert_with_pdf2docx(self):
        try:
            self.status.emit("Konwersja tekstowego PDF...")
            self.progress.emit(10)
            
            cv = Converter(self.pdf_path)
            self.progress.emit(30)
            
            cv.convert(self.output_path, start=0, end=None)
            self.progress.emit(80)
            
            cv.close()
            self.progress.emit(100)
            
            self.finished.emit(True, 
                "✅ Konwersja zakończona!\n\n"
                "PDF zawierał tekst - użyto szybkiej konwersji.\n"
                "Dokument Word jest edytowalny.")
                
        except Exception as e:
            self.log.emit(f"Błąd pdf2docx: {e}, przełączam na OCR...")
            self.convert_with_ocr()

    def convert_with_ocr(self):
        try:
            if not OCR_AVAILABLE:
                self.finished.emit(False, 
                    "Brak bibliotek OCR!\n\n"
                    "Zainstaluj:\n"
                    "pip install pytesseract pdf2image pillow python-docx\n\n"
                    "ORAZ pobierz i zainstaluj Tesseract OCR z:\n"
                    "https://github.com/UB-Mannheim/tesseract/wiki")
                return

            # Ścieżki do Tesseract i Poppler (dostosuj do swojego systemu)
            pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            poppler_path = r'C:\Program Files\poppler-25.12.0\Library\bin'
            
            self.status.emit("Konwersja PDF do obrazów...")
            self.progress.emit(10)
            
            images = convert_from_path(
                self.pdf_path, 
                dpi=self.dpi,
                fmt='png',
                poppler_path=poppler_path
            )
            
            total_pages = len(images)
            self.log.emit(f"Znaleziono {total_pages} stron")
            
            doc = docx.Document()
            doc.add_heading('Dokument przekonwertowany z PDF', 0)
            
            for i, image in enumerate(images):
                self.progress.emit(int(10 + (i / total_pages) * 80))
                self.status.emit(f"OCR strony {i+1}/{total_pages}...")
                
                temp_img = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                temp_path = temp_img.name
                temp_img.close()
                image.save(temp_path, 'PNG')
                
                img = Image.open(temp_path)
                text = pytesseract.image_to_string(
                    img, 
                    lang=self.lang,
                    config='--psm 6'
                )
                img.close()
                
                if text.strip():
                    doc.add_heading(f'Strona {i+1}', level=1)
                    paragraphs = text.split('\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                
                time.sleep(0.1)
                try:
                    os.unlink(temp_path)
                except:
                    pass
                
                if i < total_pages - 1:
                    doc.add_page_break()
            
            self.status.emit("Zapisywanie dokumentu Word...")
            self.progress.emit(95)
            
            doc.save(self.output_path)
            
            self.progress.emit(100)
            self.status.emit("Gotowe!")
            
            self.finished.emit(True, 
                f"✅ Konwersja OCR zakończona!\n\n"
                f"📄 Przetworzono: {total_pages} stron\n"
                f"📁 Zapisano: {self.output_path}\n\n"
                f"Tekst został rozpoznany i jest edytowalny!\n"
                f"⚠️ Sprawdź czy wszystkie znaki są poprawne.")
                
        except Exception as e:
            error_msg = str(e)
            if "tesseract is not installed" in error_msg:
                error_msg = "Tesseract OCR nie jest zainstalowany!\n\nPobierz z:\nhttps://github.com/UB-Mannheim/tesseract/wiki"
            elif "poppler" in error_msg.lower():
                error_msg = "Brak Popplera!\n\nDodaj ścieżkę do Popplera w PATH lub zainstaluj go."
            self.finished.emit(False, f"Błąd OCR: {error_msg}")


class ConverterTab(QWidget):
    """Zakładka konwertera PDF → Word z OCR"""

    def __init__(self, main_window=None):
        super().__init__()
        self.main_window = main_window
        self.pdf_path = None
        self.setup_ui()
        self.check_dependencies()

    def check_dependencies(self):
        status_text = ""
        
        if PDF2DOCX_AVAILABLE:
            status_text += "✅ pdf2docx (dla tekstowych PDF)\n"
        else:
            status_text += "⚠️ pdf2docx - brak (pip install pdf2docx)\n"
            
        if OCR_AVAILABLE:
            status_text += "✅ OCR (dla zeskanowanych PDF)\n"
        else:
            status_text += "❌ OCR - brak (pip install pytesseract pdf2image pillow python-docx)\n"
            
        self.status_label.setText(status_text)
        
        if OCR_AVAILABLE or PDF2DOCX_AVAILABLE:
            self.convert_btn.setEnabled(True)
        else:
            self.convert_btn.setEnabled(False)
            self.status_label.setStyleSheet("color: #f38ba8;")

    def setup_ui(self):
        main_layout = QVBoxLayout()
        main_layout.setSpacing(15)

        self.status_label = QLabel("Sprawdzanie bibliotek...")
        self.status_label.setAlignment(Qt.AlignLeft)
        self.status_label.setWordWrap(True)
        self.status_label.setStyleSheet("font-family: monospace; padding: 5px; background-color: #1e1e2e; border-radius: 6px; color: #cdd6f4;")
        main_layout.addWidget(self.status_label)

        file_group = QGroupBox("📁 Wybór pliku")
        file_group.setStyleSheet("QGroupBox { color: #89b4fa; font-weight: bold; }")
        file_layout = QHBoxLayout()

        self.file_label = QLabel("📄 Brak wybranego pliku")
        self.file_label.setStyleSheet("padding: 8px; background-color: #1e1e2e; border-radius: 5px; color: #cdd6f4;")
        file_layout.addWidget(self.file_label, 1)

        self.select_btn = QPushButton("📂 Wybierz PDF")
        self.select_btn.clicked.connect(self.select_pdf)
        self.select_btn.setStyleSheet(self._btn_style())
        file_layout.addWidget(self.select_btn)

        file_group.setLayout(file_layout)
        main_layout.addWidget(file_group)

        mode_group = QGroupBox("🎯 Tryb konwersji")
        mode_group.setStyleSheet("QGroupBox { color: #89b4fa; }")
        mode_layout = QVBoxLayout()
        
        self.auto_mode = QRadioButton("🤖 Automatyczny (zalecany)")
        self.auto_mode.setChecked(True)
        self.auto_mode.setStyleSheet("color: #cdd6f4;")
        mode_layout.addWidget(self.auto_mode)
        
        self.ocr_mode = QRadioButton("🖼️ OCR (dla zeskanowanych PDF / obrazów)")
        self.ocr_mode.setStyleSheet("color: #cdd6f4;")
        mode_layout.addWidget(self.ocr_mode)
        
        self.pdf2docx_mode = QRadioButton("📄 PDF2DOCX (tylko dla tekstowych PDF)")
        self.pdf2docx_mode.setStyleSheet("color: #cdd6f4;")
        mode_layout.addWidget(self.pdf2docx_mode)
        
        mode_group.setLayout(mode_layout)
        main_layout.addWidget(mode_group)

        ocr_group = QGroupBox("⚙️ Opcje OCR")
        ocr_group.setStyleSheet("QGroupBox { color: #89b4fa; }")
        ocr_layout = QVBoxLayout()
        
        lang_layout = QHBoxLayout()
        lang_label = QLabel("Język dokumentu:")
        lang_label.setStyleSheet("color: #cdd6f4;")
        lang_layout.addWidget(lang_label)
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(["polski (pol)", "angielski (eng)", "pol+eng", "niemiecki (deu)", "francuski (fra)"])
        self.lang_combo.setCurrentText("pol+eng")
        self.lang_combo.setStyleSheet("background-color: #1e1e2e; color: #cdd6f4; padding: 5px; border-radius: 4px;")
        lang_layout.addWidget(self.lang_combo)
        lang_layout.addStretch()
        ocr_layout.addLayout(lang_layout)
        
        dpi_layout = QHBoxLayout()
        dpi_label = QLabel("Jakość OCR (DPI):")
        dpi_label.setStyleSheet("color: #cdd6f4;")
        dpi_layout.addWidget(dpi_label)
        self.dpi_spin = QSpinBox()
        self.dpi_spin.setRange(150, 600)
        self.dpi_spin.setValue(300)
        self.dpi_spin.setStyleSheet("background-color: #1e1e2e; color: #cdd6f4; border: 1px solid #313244; border-radius: 4px;")
        dpi_layout.addWidget(self.dpi_spin)
        dpi_info = QLabel(" (300 = standard, 600 = bardzo dokładne)")
        dpi_info.setStyleSheet("color: #a6adc8;")
        dpi_layout.addWidget(dpi_info)
        dpi_layout.addStretch()
        ocr_layout.addLayout(dpi_layout)
        
        ocr_group.setLayout(ocr_layout)
        main_layout.addWidget(ocr_group)

        info_group = QGroupBox("ℹ️ Informacja")
        info_group.setStyleSheet("QGroupBox { color: #89b4fa; }")
        info_layout = QVBoxLayout()
        
        info_text = QLabel(
            "📌 **Jak to działa?**\n\n"
            "1. **PDF tekstowy** (zaznaczalny tekst) → szybka konwersja, edytowalny Word\n"
            "2. **PDF zeskanowany** (obrazy) → OCR rozpoznaje tekst, tworzy edytowalny Word\n\n"
            "✅ **Zachowuje edytowalność** - to NIE jest obraz w Wordzie!\n"
            "✅ Rozpoznaje polskie znaki (ą, ć, ę, ł, ń, ó, ś, ź, ż)\n\n"
            "⚠️ Pierwsza konwersja może trwać dłużej (ładowanie modeli OCR)"
        )
        info_text.setWordWrap(True)
        info_text.setStyleSheet("color: #a6adc8; font-size: 11px; padding: 5px;")
        info_layout.addWidget(info_text)
        
        info_group.setLayout(info_layout)
        main_layout.addWidget(info_group)

        log_group = QGroupBox("📝 Logi konwersji")
        log_group.setStyleSheet("QGroupBox { color: #89b4fa; }")
        log_layout = QVBoxLayout()
        
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(120)
        self.log_text.setStyleSheet("font-family: monospace; font-size: 10px; background-color: #1e1e2e; color: #cdd6f4; border: 1px solid #313244; border-radius: 6px;")
        log_layout.addWidget(self.log_text)
        
        log_group.setLayout(log_layout)
        main_layout.addWidget(log_group)

        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                background-color: #1e1e2e;
                border: 1px solid #313244;
                border-radius: 4px;
                text-align: center;
                color: #cdd6f4;
            }
            QProgressBar::chunk {
                background-color: #89b4fa;
                border-radius: 4px;
            }
        """)
        main_layout.addWidget(self.progress_bar)
        
        self.status_progress_label = QLabel("")
        self.status_progress_label.setVisible(False)
        self.status_progress_label.setStyleSheet("color: #a6adc8;")
        main_layout.addWidget(self.status_progress_label)

        self.convert_btn = QPushButton("🔄 KONWERTUJ DO WORD (edytowalny)")
        self.convert_btn.setEnabled(False)
        self.convert_btn.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                font-weight: bold;
                padding: 12px;
                border-radius: 8px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #2ecc71;
            }
            QPushButton:disabled {
                background-color: #45475a;
                color: #6c7086;
            }
        """)
        self.convert_btn.clicked.connect(self.convert_to_word)
        main_layout.addWidget(self.convert_btn)

        main_layout.addStretch()
        self.setLayout(main_layout)

    def _btn_style(self):
        return """
            QPushButton {
                background-color: #313244;
                color: #cdd6f4;
                border: none;
                border-radius: 6px;
                padding: 8px 16px;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #89b4fa;
                color: #1a1a2e;
            }
        """

    def select_pdf(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Wybierz plik PDF", "", "Pliki PDF (*.pdf)"
        )

        if file_path:
            self.pdf_path = file_path
            self.file_label.setText(f"📄 {os.path.basename(file_path)}")
            self.convert_btn.setEnabled(True)
            self.log_text.clear()
            self.log_text.append(f"✅ Wybrano: {os.path.basename(file_path)}")
            self._set_status(f"Wybrano plik: {os.path.basename(file_path)}")

    def convert_to_word(self):
        if not self.pdf_path:
            QMessageBox.warning(self, "Uwaga", "Najpierw wybierz plik PDF!")
            return

        default_name = os.path.splitext(os.path.basename(self.pdf_path))[0] + "_konwertowany.docx"
        output_path, _ = QFileDialog.getSaveFileName(
            self, "Zapisz jako Word", default_name, 
            "Pliki Word (*.docx);;Wszystkie pliki (*.*)"
        )

        if output_path:
            if self.auto_mode.isChecked():
                mode = "auto"
            elif self.ocr_mode.isChecked():
                mode = "ocr"
            else:
                mode = "pdf2docx"
            
            lang_map = {
                "polski (pol)": "pol",
                "angielski (eng)": "eng",
                "pol+eng": "pol+eng",
                "niemiecki (deu)": "deu",
                "francuski (fra)": "fra"
            }
            lang = lang_map[self.lang_combo.currentText()]
            
            self.progress_bar.setVisible(True)
            self.status_progress_label.setVisible(True)
            self.progress_bar.setValue(0)
            self.convert_btn.setEnabled(False)
            self.select_btn.setEnabled(False)
            self.log_text.clear()

            self.convert_thread = OCRConvertThread(
                self.pdf_path, 
                output_path,
                mode=mode,
                dpi=self.dpi_spin.value(),
                lang=lang
            )
            self.convert_thread.progress.connect(self.progress_bar.setValue)
            self.convert_thread.status.connect(self.status_progress_label.setText)
            self.convert_thread.log.connect(self.log_text.append)
            self.convert_thread.finished.connect(self.on_convert_finished)
            self.convert_thread.start()

    def on_convert_finished(self, success, message):
        self.progress_bar.setVisible(False)
        self.status_progress_label.setVisible(False)
        self.convert_btn.setEnabled(True)
        self.select_btn.setEnabled(True)

        if success:
            QMessageBox.information(self, "✅ Sukces!", message)
            self._set_status("Konwersja zakończona pomyślnie!")
        else:
            QMessageBox.critical(self, "❌ Błąd!", message)
            self._set_status("Błąd konwersji!")

    def _set_status(self, text):
        if self.main_window and hasattr(self.main_window, 'status_label'):
            self.main_window.status_label.setText(text)